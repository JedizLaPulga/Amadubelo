"""
File Converters
Handles PDF, DOCX, Image, and Text conversions.
"""

import os
from pathlib import Path
from typing import List, Optional, Callable
from PIL import Image
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader


def pdf_to_images(
    pdf_path: str,
    output_folder: str,
    format: str = "png",
    dpi: int = 200,
    progress_callback: Optional[Callable] = None
) -> List[str]:
    """
    Convert PDF pages to images.
    
    Args:
        pdf_path: Path to PDF file
        output_folder: Folder to save images
        format: Output format (png, jpg)
        dpi: Resolution (default 200)
        progress_callback: Optional callback(current, total)
        
    Returns:
        List of output image paths
    """
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Get base name
    base_name = Path(pdf_path).stem
    
    # Try PyMuPDF first (doesn't require Poppler)
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(pdf_path)
        output_files = []
        total = len(doc)
        
        # Calculate zoom based on DPI (default PDF DPI is 72)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=matrix)
            output_path = os.path.join(output_folder, f"{base_name}_page_{i+1}.{format}")
            pix.save(output_path)
            output_files.append(output_path)
            
            if progress_callback:
                progress_callback(i + 1, total)
        
        doc.close()
        return output_files
        
    except ImportError:
        pass  # PyMuPDF not installed, try pdf2image
    
    # Fallback to pdf2image (requires Poppler)
    try:
        from pdf2image import convert_from_path
        from pdf2image.exceptions import PDFInfoNotInstalledError
        
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
        except PDFInfoNotInstalledError:
            raise RuntimeError(
                "Poppler is not installed!\n\n"
                "To use PDF to Image conversion, please install PyMuPDF:\n"
                "   pip install pymupdf\n\n"
                "Or install Poppler for Windows:\n"
                "   1. Download from: https://github.com/oschwartz10612/poppler-windows/releases\n"
                "   2. Extract to C:\\Program Files\\poppler\n"
                "   3. Add C:\\Program Files\\poppler\\bin to your PATH"
            )
        except Exception as e:
            if "poppler" in str(e).lower() or "Unable to get page count" in str(e):
                raise RuntimeError(
                    "Poppler is not installed!\n\n"
                    "To use PDF to Image conversion, please install PyMuPDF:\n"
                    "   pip install pymupdf\n\n"
                    "Or install Poppler for Windows:\n"
                    "   1. Download from: https://github.com/oschwartz10612/poppler-windows/releases\n"
                    "   2. Extract to C:\\Program Files\\poppler\n"
                    "   3. Add C:\\Program Files\\poppler\\bin to your PATH"
                )
            raise
        
        output_files = []
        total = len(images)
        
        for i, image in enumerate(images):
            output_path = os.path.join(output_folder, f"{base_name}_page_{i+1}.{format}")
            image.save(output_path, format.upper())
            output_files.append(output_path)
            
            if progress_callback:
                progress_callback(i + 1, total)
                
        return output_files
        
    except ImportError:
        raise RuntimeError(
            "No PDF rendering library installed!\n\n"
            "Please install PyMuPDF:\n"
            "   pip install pymupdf"
        )


def images_to_pdf(
    image_paths: List[str],
    output_path: str,
    page_size: str = "A4",
    progress_callback: Optional[Callable] = None
) -> str:
    """
    Combine multiple images into a single PDF.
    
    Args:
        image_paths: List of image file paths
        output_path: Output PDF path
        page_size: Page size (A4, letter)
        progress_callback: Optional callback(current, total)
        
    Returns:
        Output PDF path
    """
    if not image_paths:
        raise ValueError("No images provided")
    
    # Determine page size
    sizes = {"A4": A4, "letter": letter}
    size = sizes.get(page_size, A4)
    
    # Create PDF
    c = canvas.Canvas(output_path, pagesize=size)
    page_width, page_height = size
    
    total = len(image_paths)
    
    for i, img_path in enumerate(image_paths):
        img = Image.open(img_path)
        img_width, img_height = img.size
        
        # Calculate scaling to fit page
        width_ratio = page_width / img_width
        height_ratio = page_height / img_height
        scale = min(width_ratio, height_ratio) * 0.95  # 5% margin
        
        new_width = img_width * scale
        new_height = img_height * scale
        
        # Center on page
        x = (page_width - new_width) / 2
        y = (page_height - new_height) / 2
        
        # Draw image
        c.drawImage(ImageReader(img), x, y, width=new_width, height=new_height)
        
        if i < total - 1:
            c.showPage()
            
        if progress_callback:
            progress_callback(i + 1, total)
    
    c.save()
    return output_path


def text_to_pdf(
    text_path: str,
    output_path: str,
    font_size: int = 12,
    font_name: str = "Helvetica"
) -> str:
    """
    Convert text file to PDF.
    
    Args:
        text_path: Path to text file
        output_path: Output PDF path
        font_size: Font size
        font_name: Font name
        
    Returns:
        Output PDF path
    """
    # Read text content
    with open(text_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create PDF
    c = canvas.Canvas(output_path, pagesize=A4)
    page_width, page_height = A4
    
    # Set font
    c.setFont(font_name, font_size)
    
    # Calculate line height
    line_height = font_size * 1.5
    margin = 50
    y = page_height - margin
    max_width = page_width - (2 * margin)
    
    # Split into lines
    lines = content.split('\n')
    
    for line in lines:
        # Handle empty lines
        if not line.strip():
            y -= line_height
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = page_height - margin
            continue
            
        # Word wrap
        words = line.split()
        current_line = ""
        
        for word in words:
            test_line = f"{current_line} {word}".strip()
            text_width = c.stringWidth(test_line, font_name, font_size)
            
            if text_width < max_width:
                current_line = test_line
            else:
                # Draw current line
                c.drawString(margin, y, current_line)
                y -= line_height
                
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = page_height - margin
                    
                current_line = word
        
        # Draw remaining text
        if current_line:
            c.drawString(margin, y, current_line)
            y -= line_height
            
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = page_height - margin
    
    c.save()
    return output_path


def docx_to_pdf(
    docx_path: str,
    output_path: str,
    progress_callback: Optional[Callable] = None
) -> str:
    """
    Convert DOCX to PDF.
    
    Note: This is a simplified conversion that extracts text.
    For full formatting, consider using LibreOffice or MS Word COM.
    
    Args:
        docx_path: Path to DOCX file
        output_path: Output PDF path
        progress_callback: Optional callback
        
    Returns:
        Output PDF path
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    # Read DOCX
    doc = Document(docx_path)
    
    # Create PDF
    c = canvas.Canvas(output_path, pagesize=A4)
    page_width, page_height = A4
    
    margin = 50
    y = page_height - margin
    line_height = 14
    
    total_paragraphs = len(doc.paragraphs)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            y -= line_height
            if y < margin:
                c.showPage()
                y = page_height - margin
            continue
        
        # Simple style handling
        font_size = 12
        font_name = "Helvetica"
        
        if para.style.name.startswith("Heading"):
            font_size = 16
            font_name = "Helvetica-Bold"
        
        c.setFont(font_name, font_size)
        
        # Word wrap
        words = text.split()
        current_line = ""
        max_width = page_width - (2 * margin)
        
        for word in words:
            test_line = f"{current_line} {word}".strip()
            text_width = c.stringWidth(test_line, font_name, font_size)
            
            if text_width < max_width:
                current_line = test_line
            else:
                c.drawString(margin, y, current_line)
                y -= line_height * 1.2
                
                if y < margin:
                    c.showPage()
                    y = page_height - margin
                    c.setFont(font_name, font_size)
                    
                current_line = word
        
        if current_line:
            c.drawString(margin, y, current_line)
            y -= line_height * 1.5
            
            if y < margin:
                c.showPage()
                y = page_height - margin
        
        if progress_callback:
            progress_callback(i + 1, total_paragraphs)
    
    c.save()
    return output_path
